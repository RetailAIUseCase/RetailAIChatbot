"""
Document parsing utilities for different file formats
"""
import os
import re
import json
import hashlib
from typing import List, Dict, Any, Optional
import logging
from io import BytesIO
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)

class DocumentParser:
    """Base class for document parsers with hash-based duplicate checking"""
    
    @staticmethod
    def generate_content_hash(content: str, project_context: str = "") -> str:
        # Include project context to reduce cross-project collisions
        content_with_context = f"{project_context}:{content}" if project_context else content
        return hashlib.sha256(content_with_context.encode('utf-8')).hexdigest()
    
    @staticmethod
    async def check_embedding_by_content_hash(
        connection, 
        user_id: int, 
        project_id: str, 
        content_hash: str, 
        embedding_table: str
    ) -> Optional[Dict[str, Any]]:
        """Check if embedding already exists for this content hash anywhere in the project"""
        query = f"""
        SELECT id, content, embedding, metadata, document_id, created_at
        FROM {embedding_table} 
        WHERE user_id = $1 AND project_id = $2 
              AND metadata->>'content_hash' = $3
        LIMIT 1
        """
        try:
            row = await connection.fetchrow(query, user_id, project_id, content_hash)
            return dict(row) if row else None
        except Exception as e:
            logger.error(f"Error checking embedding by content hash: {e}")
            return None
        
    @staticmethod
    async def check_existing_metadata_embedding(
        connection, 
        table_name: str,
        # document_id: str, 
        user_id: int, 
        project_id: str, 
        content_type: str
    ) -> Optional[Dict[str, Any]]:
        """Check if embedding already exists for this content type and document"""
        
        query = """
        SELECT content, embedding, metadata, content_type
        FROM metadata_embeddings 
        WHERE user_id = $1 AND project_id = $2 
                AND table_name = $3 AND content_type = $4
        """
        # params = [document_id, user_id, project_id, table_name, content_type]
        params = [user_id, project_id, table_name, content_type]
        try:
            row = await connection.fetchrow(query, *params)
            return dict(row) if row else None
        except Exception as e:
            logger.error(f"Error checking existing embedding: {e}")
            return None
    
    @staticmethod
    def should_update_embedding(existing_record: Optional[Dict], new_content: str, new_embedding: List[float]) -> tuple[bool, str]:
        """
        Check if embedding should be updated
        Returns: (should_update: bool, action: str)
        """
        new_content_hash = DocumentParser.generate_content_hash(new_content)
        
        if not existing_record:
            return True, "INSERT"
        
        # Get existing content hash
        existing_metadata = existing_record.get('metadata', {})
        if isinstance(existing_metadata, str):
            try:
                existing_metadata = json.loads(existing_metadata)
            except:
                existing_metadata = {}
        
        existing_content_hash = existing_metadata.get('content_hash')
        
        # If no existing hash or hashes differ, update
        if not existing_content_hash or existing_content_hash != new_content_hash:
            return True, "UPDATE"
        
        # Check if embeddings are different (with small tolerance for floating point)
        existing_embedding = existing_record.get('embedding', [])
        if len(existing_embedding) != len(new_embedding):
            return True, "UPDATE"
        
        # Compare embeddings with tolerance
        for i, (old_val, new_val) in enumerate(zip(existing_embedding, new_embedding)):
            if abs(old_val - new_val) > 1e-8:  # Small tolerance for float comparison
                return True, "UPDATE"
        
        return False, "SKIP"

class MetadataParser(DocumentParser):
    """Parser for metadata documents (JSON and DOCX)"""
    
    @staticmethod
    def is_json_content(text: str) -> bool:
        """Check if the content is JSON format"""
        try:
            json.loads(text.strip())
            return True
        except (json.JSONDecodeError, ValueError):
            return False
    
    @classmethod
    def parse_json_metadata(cls, text: str) -> List[Dict]:
        """Parse metadata from JSON format"""
        try:
            data = json.loads(text.strip())
            tables = []
            
            if isinstance(data, dict):
                if "tables" in data:
                    for table in data["tables"]:
                        tables.append(cls._normalize_table_structure(table))
                elif "schema" in data:
                    for table in data["schema"].get("tables", []):
                        tables.append(cls._normalize_table_structure(table))
                else:
                    tables.append(cls._normalize_table_structure(data))
            elif isinstance(data, list):
                for table in data:
                    tables.append(cls._normalize_table_structure(table))
            
            return tables
            
        except Exception as e:
            logger.error(f"Error parsing JSON metadata: {e}")
            return []
    
    @staticmethod
    def _normalize_table_structure(table_data: Dict) -> Dict:
        """Normalize table structure from JSON to consistent format"""
        normalized = {
            "name": table_data.get("name") or table_data.get("table_name") or table_data.get("tableName", "unknown_table"),
            "purpose": table_data.get("purpose") or table_data.get("description") or table_data.get("comment", ""),
            "columns": [],
            "primary_key": "",
            "foreign_keys": "",
            # Enhanced fields from your rich JSON format
            "grain": table_data.get("grain", ""),
            "entities": table_data.get("entities", ""),
            "key_columns_summary": table_data.get("key_columns_summary", ""),
            "llm_notes": table_data.get("llm_notes", ""),
            "synonyms": table_data.get("synonyms", "")
        }
        
        # Handle columns (existing logic works fine)
        columns_data = table_data.get("columns") or table_data.get("fields") or []
        for col in columns_data:
            if isinstance(col, dict):
                normalized["columns"].append({
                    "name": col.get("name") or col.get("field_name") or col.get("column_name", "unknown_column"),
                    "type": col.get("type") or col.get("data_type") or col.get("dataType", "UNKNOWN"),
                    "description": col.get("description") or col.get("comment") or col.get("desc", "")
                })
            elif isinstance(col, str):
                normalized["columns"].append({
                    "name": col,
                    "type": "UNKNOWN", 
                    "description": ""
                })
        
        # Handle primary key
        pk_data = table_data.get("primary_key") or table_data.get("primaryKey") or table_data.get("pk")
        if pk_data:
            if isinstance(pk_data, list):
                normalized["primary_key"] = ", ".join(pk_data)
            else:
                normalized["primary_key"] = str(pk_data)
        
        # Handle foreign keys
        fk_data = table_data.get("foreign_keys") or table_data.get("foreignKeys") or table_data.get("fk") or []
        if fk_data:
            if isinstance(fk_data, list):
                normalized["foreign_keys"] = "; ".join([str(fk) for fk in fk_data])
            else:
                normalized["foreign_keys"] = str(fk_data)
        
        return normalized
    
    @classmethod
    def parse_docx_metadata(cls, text: str) -> List[Dict]:
        """Parse metadata from DOCX format"""
        tables = []
        current_table = None
        current_section = None
        
        # Regex patterns
        table_re = re.compile(r"^(?:\d+\.\s*)?Table(?:\s+Name)?\s*[-–]\s*([A-Za-z0-9_]+)\s*$", re.IGNORECASE | re.MULTILINE)
        fields_hdr_re = re.compile(r"^Fields?\s*[-–]?\s*$", re.IGNORECASE)
        pk_hdr_re = re.compile(r"^Primary\s+Key[s]?\s*[-–]?\s*$", re.IGNORECASE)
        fk_hdr_re = re.compile(r"^Foreign\s+Key[s]?\s*[-–]?\s*$", re.IGNORECASE)
        field_re = re.compile(r"^([A-Za-z_][A-Za-z0-9_]*)\s+([A-Za-z]+(?:\(\s*\d+(?:\s*,\s*\d+)?\s*\))?)\s+\[Comment\s*[-–]\s*(.*?)\]\s*$", re.IGNORECASE)
        purpose_re = re.compile(r"\[Comment\s*[-–]\s*(.*?)\]", re.IGNORECASE)
        
        def flush_current():
            if current_table:
                pk_list = current_table.pop("primary_keys", [])
                fk_list = current_table.pop("foreign_keys_list", [])
                
                current_table["primary_key"] = ", ".join([p for p in pk_list if p])
                current_table["foreign_keys"] = "; ".join([fk for fk in fk_list if fk])
                current_table["name"] = current_table.pop("table_name", "unknown_table")
                
                tables.append(current_table)
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # New table header
            m_table = table_re.match(line)
            if m_table:
                flush_current()
                table_name = m_table.group(1)
                current_table = {
                    "table_name": table_name,
                    "purpose": "",
                    "columns": [],
                    "primary_keys": [],
                    "foreign_keys_list": []
                }
                current_section = "table_start"
                continue
            
            if not current_table:
                continue
            
            # Purpose
            if current_section == "table_start":
                m_purpose = purpose_re.search(line)
                if m_purpose:
                    current_table["purpose"] = m_purpose.group(1).strip()
                    current_section = None
                    continue
            
            # Section headers
            if fields_hdr_re.match(line):
                current_section = "fields"
                continue
            if pk_hdr_re.match(line):
                current_section = "primary_key"
                continue
            if fk_hdr_re.match(line):
                current_section = "foreign_key"
                continue
            
            # Content parsing
            if current_section == "fields":
                mf = field_re.match(line)
                if mf:
                    col_name, col_type, col_desc = mf.groups()
                    current_table["columns"].append({
                        "name": col_name.strip(),
                        "type": col_type.strip().upper(),
                        "description": col_desc.strip()
                    })
                continue
            
            if current_section == "primary_key":
                if not line.lower().startswith("primary"):
                    parts = [p.strip() for p in line.split(",") if p.strip()]
                    current_table["primary_keys"].extend(parts)
                continue
            
            if current_section == "foreign_key":
                if not line.lower().startswith("foreign"):
                    current_table["foreign_keys_list"].append(line.strip())
                continue
        
        flush_current()
        return tables
    
    # @staticmethod
    # def extract_relationships(table: Dict) -> List[str]:
    #     """Extract relationships from table definition"""
    #     relationships = []
    #     fk_string = table.get('foreign_keys', '')
        
    #     if fk_string:
    #         for fk in fk_string.split(';'):
    #             fk = fk.strip()
    #             if not fk:
    #                 continue
                
    #             if ' references ' in fk.lower():
    #                 try:
    #                     parts = re.split(r'\s+references\s+', fk, flags=re.IGNORECASE)
    #                     if len(parts) == 2:
    #                         local_cols = parts[0].strip()
    #                         ref_table = parts[1].split('(')[0].strip()
    #                         relationships.append(f"Relates to {ref_table} via {local_cols}")
    #                 except Exception as e:
    #                     logger.error(f"Couldn't parse FK '{fk}': {str(e)}")
    #                     relationships.append(f"Complex relationship: {fk}")
    #             else:
    #                 relationships.append(fk)  # Include as-is for rich descriptions like "(Referenced by many tables)"
        
    #     return relationships
    
    @classmethod
    async def create_embeddings_with_dedup(
        cls, 
        connection,
        tables: List[Dict], 
        document_id: str, 
        user_id: int, 
        project_id: str,
        get_embedding_func
    ) -> Dict[str, int]:
        """Create embeddings with deduplication with content-hash based deduplication"""
        stats = {"inserted": 0, "updated": 0, "skipped": 0,  "reused":0}
        
        for table in tables:
            # Process each view type (table, column, relationship)
            views = cls.extract_hierarchical_views(table)
            
            for view in views:
                try:
                    content = view["content"]
                    # content_hash = cls.generate_content_hash(content)
                    if view["content_type"] == "relationship":
                        project_context = f"project_{project_id}_table_{table['name']}"
                        content_hash = cls.generate_content_hash(content, project_context)
                    else:
                        content_hash = cls.generate_content_hash(content)

                    # Check if content already exist in the project
                    existing_by_hash = await cls.check_embedding_by_content_hash(
                        connection,
                        user_id,
                        project_id,
                        content_hash,
                        "metadata_embeddings"
                    )
                    if existing_by_hash:
                        # Update document_id to reflect latest upload
                        await cls._refresh_document_id_for_hash(
                            connection, document_id, user_id, project_id, content_hash
                        )
                        stats["reused"] += 1
                        logger.info(f"Reused existing embedding for {table['name']}-{view['content_type']} (content unchanged)")
                        continue

                    # Check if embedding exists for the document specific
                    existing_for_doc = await cls.check_existing_metadata_embedding(
                        connection,
                        table["name"],
                        user_id, 
                        project_id,
                        view["content_type"]
                    )
                    
                    # Generate new embedding
                    new_embedding = await get_embedding_func(content)

                    view_metadata = view.get("metadata",{})
                    view_metadata["project_id"] = project_id
                    view_metadata["content_hash"] = content_hash
                    
                    if existing_for_doc:
                        # Update exisiting embedding with new content
                        await cls._update_metadata_embedding(
                            connection, document_id, user_id, project_id, 
                            table["name"], view["content_type"], content,
                            new_embedding, view_metadata
                        )
                        stats["updated"]+=1
                        logger.info(f"Updated embedding for {table['name']}-{view['content_type']} (content changed)")
                    else:
                        # Insert new embedding
                        # Insert completely new embedding
                        await cls._insert_metadata_embedding(
                            connection, document_id, user_id, project_id,
                            table["name"], view["content_type"], content,
                            new_embedding, view_metadata
                        )
                        stats["inserted"] += 1
                        logger.info(f"Created new embedding for {table['name']}-{view['content_type']}")
                
                except Exception as e:
                    logger.error(f"Error processing embedding for {table['name']}-{view['content_type']}: {e}")
                    continue
                    # Check if update needed
                #     should_update, action = cls.should_update_embedding(
                #         existing, 
                #         view["content"], 
                #         new_embedding
                #     )
                    
                #     if should_update:
                #         # Add content hash to metadata
                #         view_metadata = view.get("metadata", {})
                #         view_metadata["content_hash"] = cls.generate_content_hash(view["content"])
                        
                #         if action == "INSERT":
                #             await cls._insert_metadata_embedding(
                #                 connection, document_id, user_id, project_id,
                #                 table["name"], view["content_type"], view["content"],
                #                 new_embedding, view_metadata
                #             )
                #             stats["inserted"] += 1
                #         else:  # UPDATE
                #             await cls._update_metadata_embedding(
                #                 connection, document_id, user_id, project_id,
                #                 table["name"], view["content_type"], view["content"],
                #                 new_embedding, view_metadata
                #             )
                #             stats["updated"] += 1
                #     else:
                #         stats["skipped"] += 1
                #         logger.debug(f"Skipped unchanged embedding for {table['name']}-{view['content_type']}")
                
                # except Exception as e:
                #     logger.error(f"Error processing embedding for {table['name']}-{view['content_type']}: {e}")
                #     continue
        
        return stats
    
    @staticmethod
    async def _insert_metadata_embedding(
        connection, document_id, user_id, project_id, 
        table_name, content_type, content, embedding, metadata
    ):
        """Insert new metadata embedding"""
        insert_query = """
        INSERT INTO metadata_embeddings
        (document_id, project_id, user_id, table_name, content_type, content, embedding, metadata)
        VALUES ($1, $2, $3, $4, $5, $6, $7::vector, $8)
        """
        # Debug logging
        logger.debug(f"Inserting metadata embedding: table={table_name}, type={content_type}")
        logger.debug(f"Embedding type: {type(embedding)}, length: {len(embedding) if isinstance(embedding, list) else 'N/A'}")
        await connection.execute(
            insert_query, document_id, project_id, user_id,
            table_name, content_type, content, embedding, json.dumps(metadata)
        )
    @staticmethod
    async def _refresh_document_id_for_hash(
        connection, document_id: str, user_id: int, project_id: str, content_hash: str
    ):
        """Update the document_id for an existing embedding with matching content hash"""
        query = """
        UPDATE metadata_embeddings
        SET document_id = $1,
            created_at = CURRENT_TIMESTAMP
        WHERE user_id = $2 AND project_id = $3
            AND metadata->>'content_hash' = $4
        """
        try:
            await connection.execute(query, document_id, user_id, project_id, content_hash)
            logger.info(f"Refreshed document_id for reused embedding with hash {content_hash[:10]}...")
        except Exception as e:
            logger.error(f"Error refreshing document_id for hash {content_hash[:10]}: {e}")

    @staticmethod
    async def _update_metadata_embedding(
        connection, document_id, user_id, project_id,
        table_name, content_type, content, embedding, metadata
    ):
        """Update existing metadata embedding"""
        update_query = """
        UPDATE metadata_embeddings 
        SET content = $1, embedding = $2::vector, metadata = $3, document_id = $4, created_at = CURRENT_TIMESTAMP
        WHERE user_id = $5 AND project_id = $6 
              AND table_name = $7 AND content_type = $8
        """
        # Debug logging
        logger.debug(f"Updating metadata embedding: table={table_name}, type={content_type}")
        logger.debug(f"Embedding type: {type(embedding)}, length: {len(embedding) if isinstance(embedding, list) else 'N/A'}")
        await connection.execute(
            update_query, content, embedding, json.dumps(metadata),
            document_id, user_id, project_id, table_name, content_type
        )
    
    @staticmethod
    def extract_hierarchical_views(table: Dict[str, Any]) -> List[Dict]:
        """Generate enhanced views for each table with rich metadata"""
        views = []
        
        # 1. Enhanced Table Summary View
        table_summary = f"""TABLE: {table['name']}
            PURPOSE: {table.get('purpose', 'N/A')}
            GRAIN: {table.get('grain', 'N/A')}
            PRIMARY KEY: {table.get('primary_key', 'N/A')}
            FOREIGN KEYS: {table.get('foreign_keys', 'N/A')}
            COLUMN COUNT: {len(table.get('columns', []))}
            SYNONYMS: {table.get('synonyms', 'N/A')}"""

        views.append({
            "content_type": "table",
            "content": table_summary,
            "metadata": {
                "purpose": table.get('purpose'),
                "grain": table.get('grain'),
                "primary_key": table.get('primary_key'),
                "foreign_keys": table.get('foreign_keys'),
                "synonyms": table.get('synonyms')
            }
        })
        
        # 2. Enhanced Column Descriptions View
        # columns_by_category = {
        #     "identifier": [],
        #     "measure": [],
        #     "attribute": [],
        #     "metadata": []
        # }
        
        # for col in table.get('columns', []):
        #     if col['name'] in [pk.strip() for pk in table.get('primary_key', '').split(',')]:
        #         columns_by_category["identifier"].append(col)
        #     elif any(t in col['type'].lower() for t in ['int', 'num', 'float', 'decimal']):
        #         columns_by_category["measure"].append(col)
        #     else:
        #         columns_by_category["attribute"].append(col)
        
        # column_view = "COLUMNS:\n"
        # if table.get('key_columns_summary'):
        #     column_view += f"KEY COLUMNS SUMMARY: {table['key_columns_summary']}\n\n"
        
        # for category, cols in columns_by_category.items():
        #     if cols:
        #         column_view += f"\n{category.upper()}:\n"
        #         column_view += "\n".join(
        #             f"- {col['name']} ({col['type']}): {col['description']}"
        #             for col in cols
        #         )
        
        # views.append({
        #     "content_type": "column",
        #     "content": column_view,
        #     "metadata": {
        #         "column_count": len(table.get('columns', [])),
        #         "categories": list(columns_by_category.keys()),
        #         "key_columns_summary": table.get('key_columns_summary')
        #     }
        # })
        # 2. Structured Column Information - Exact names and types only
        column_view = "AVAILABLE_COLUMNS:\n"
        for col in table.get('columns', []):
            column_view += f"- {col['name']} ({col['type']}): {col['description']}\n"
        
        views.append({
            "content_type": "column", 
            "content": column_view,
            "metadata": {
                "column_names": [col['name'] for col in table.get('columns', [])],
                "column_types": {col['name']: col['type'] for col in table.get('columns', [])}
            }
        })
        
        # 3. Enhanced Relationships & Context View
        # relationships = []
        # fk_string = table.get('foreign_keys', '')
        # if fk_string:
        #     for fk in fk_string.split(';'):
        #         fk = fk.strip()
        #         if not fk:
        #             continue
                    
        #         if ' references ' in fk.lower():
        #             try:
        #                 parts = re.split(r'\s+references\s+', fk, flags=re.IGNORECASE)
        #                 if len(parts) == 2:
        #                     local_cols = parts[0].strip()
        #                     ref_table = parts[1].split('(')[0].strip()
        #                     relationships.append(f"Relates to {ref_table} via {local_cols}")
        #             except Exception as e:
        #                 logger.error(f"Couldn't parse FK '{fk}': {str(e)}")
        #                 relationships.append(f"Complex relationship: {fk}")
        #         else:
        #             relationships.append(fk)  # Include as-is for rich descriptions like "(Referenced by many tables)"
        
        # relationship_view = "RELATIONSHIPS:\n"
        # if relationships:
        #     relationship_view += "\n".join(relationships)
        # else:
        #     relationship_view += "No explicit relationships"
        
        # if table.get('llm_notes'):
        #     relationship_view += f"\n\nCONTEXT NOTES:\n{table['llm_notes']}"
        
        # views.append({
        #     "content_type": "relationship",
        #     "content": relationship_view,
        #     "metadata": {
        #         "relationship_count": len(relationships),
        #         "target_tables": list(set(
        #             r.split('via')[0].strip().replace('Relates to ', '')
        #             for r in relationships
        #             if 'via' in r)),
        #         "llm_notes": table.get('llm_notes'),
        #         "has_context": bool(table.get('llm_notes'))
        #     }
        # })
        
        # return views
        relationship_view = "TABLE_RELATIONSHIPS:\n"
        fk_string = table.get('foreign_keys', '')
        if fk_string:
            relationship_view += fk_string
        else:
            relationship_view += "No foreign key relationships"
        
        views.append({
            "content_type": "relationship",
            "content": relationship_view,
            "metadata": {
                "has_foreign_keys": bool(fk_string),
                "foreign_key_string": fk_string
            }
        })
        return views

class BusinessLogicParser(DocumentParser):
    """Parser for business logic documents"""
    @staticmethod
    def extract_business_rules(text: str) -> List[str]:
        """
        Universal business rules extractor - handles all numbering formats:
        - 1. Rule text
        - 1) Rule text  
        - 1: Rule text
        - Rule 1: Rule text
        - Rule1: Rule text
        - **Rule 1:** Rule text (with bold)
        - Policy 1. Rule text
        - etc.
        """
        
        rules = []
        
        # Normalize text
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Multiple approaches for maximum flexibility
        
        # Method 1: Universal pattern using finditer
        # This captures most common patterns
        universal_patterns = [
            # Bold markdown format: **Rule 1:**
            r'(?:^|\n)\s*\*\*\s*(?:Rule|Policy|Section|Item|Point|Step)?\s*(\d+)\s*:\s*\*\*\s*(.+?)(?=(?:\n\s*\*\*\s*(?:Rule|Policy|Section|Item|Point|Step)?\s*\d+\s*:\s*\*\*|\Z))',
            
            # Standard formats: Rule 1:, 1., 1), etc.
            r'(?:^|\n)\s*(?:(?:Rule|Policy|Section|Item|Point|Step)\s*)?(\d+)[\.\)\:]\s+(.+?)(?=(?:\n\s*(?:(?:Rule|Policy|Section|Item|Point|Step)\s*)?\d+[\.\)\:]|\Z))',
            
            # Compact format: Rule1:, Policy5.
            r'(?:^|\n)\s*(?:Rule|Policy|Section|Item|Point|Step)(\d+)[\.\:]\s+(.+?)(?=(?:\n\s*(?:Rule|Policy|Section|Item|Point|Step)\d+[\.\:]|\Z))',
        ]
        for pattern in universal_patterns:
            matches = list(re.finditer(pattern, text, flags=re.IGNORECASE | re.DOTALL))
            
            if matches:
                temp_rules = []
                for match in matches:
                    rule_number = match.group(1)
                    content = match.group(2).strip()
                    # Clean content
                    content = re.sub(r'\s+', ' ', content)
                    content = re.sub(r'\n+', ' ', content)
                    
                    if content and len(content) > 10:
                        formatted_rule = f"Rule {rule_number}: {content}"
                        temp_rules.append(formatted_rule)
                
                # If this pattern found good results, use them
                if len(temp_rules) >= 2:
                    return temp_rules
        
        # Method 2: Fallback - Split and reconstruct approach
        # Handle cases where patterns above don't work
        
        # First, insert artificial breaks before likely rule starts
        text_with_breaks = text
        
        # Add breaks before various numbering patterns
        break_patterns = [
            r'(\S)\s*(\*\*\s*(?:Rule|Policy|Section|Item|Point|Step)?\s*\d+\s*:\s*\*\*)',  # **Rule 1:**
            r'(\S)\s*((?:Rule|Policy|Section|Item|Point|Step)\s*\d+[\.\:\)])',  # Rule 1:
            r'(\S)\s*(\d+[\.\)\:])\s+([A-Z])',  # 1. Text or 1) Text
        ]
        
        for pattern in break_patterns:
            text_with_breaks = re.sub(pattern, r'\1\n\n\2', text_with_breaks, flags=re.IGNORECASE)
        
        # Clean up excessive newlines
        text_with_breaks = re.sub(r'\n{3,}', '\n\n', text_with_breaks)
        
        # Split and process
        paragraphs = text_with_breaks.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Try to match various rule formats
            rule_match = None
            
            # Try different patterns
            patterns_to_try = [
                r'^\*\*\s*(?:Rule|Policy|Section|Item|Point|Step)?\s*(\d+)\s*:\s*\*\*\s*(.+)',  # **Rule 1:**
                r'^(?:Rule|Policy|Section|Item|Point|Step)\s*(\d+)[\.\:\)]\s*(.+)',  # Rule 1:
                r'^(?:Rule|Policy|Section|Item|Point|Step)(\d+)[\.\:]\s*(.+)',  # Rule1:
                r'^(\d+)[\.\)\:]\s+(.+)',  # 1. or 1) or 1:
            ]
            
            for pattern in patterns_to_try:
                match = re.match(pattern, para, flags=re.IGNORECASE | re.DOTALL)
                if match:
                    rule_match = match
                    break
            
            if rule_match:
                rule_number = rule_match.group(1)
                content = rule_match.group(2).strip()
                
                # Clean content
                content = re.sub(r'\s+', ' ', content)
                content = re.sub(r'\n+', ' ', content)
                
                if content and len(content) > 15:
                    formatted_rule = f"Rule {rule_number}: {content}"
                    rules.append(formatted_rule)
        
        # Method 3: Final fallback - paragraph-based with number detection
        if not rules:
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                # Check if paragraph starts with any number pattern
                if re.match(r'^\s*(?:\*\*\s*)?(?:(?:Rule|Policy|Section|Item|Point|Step)\s*)?\d+[\.\)\:]', para, re.IGNORECASE):
                    rules.append(para)
        return rules



    # @staticmethod
    # def extract_business_rules(text: str) -> List[str]:
    #     """Extract numbered business rules from text"""
    #     rule_patterns = [
    #         r'\n\d+\.\s+',  # "1. Rule text"
    #         r'\n\d+\)\s+',  # "1) Rule text"
    #         r'\nRule\s+\d+[:\.]',  # "Rule 1: text"
    #         r'\nPolicy\s+\d+[:\.]',  # "Policy 1: text"
    #     ]
        
    #     for pattern in rule_patterns:
    #         parts = re.split(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
    #         print("Parts",parts,'\n')
    #         if len(parts) > 1:
    #             rules = [part.strip() for part in parts[1:] if part.strip()]
    #             print(rules,'\n')
    #             if rules:
    #                 return rules
        
    #     # Fallback to paragraphs
    #     paragraphs = text.split('\n\n')
    #     print(paragraphs,'\n')
    #     return [p.strip() for p in paragraphs if p.strip() and len(p.strip()) > 50]
    
    @classmethod
    async def create_embeddings_with_dedup(
        cls,
        connection,
        rules: List[str],
        document_id: str,
        user_id: int,
        project_id: str,
        get_embedding_func
    ) -> Dict[str, int]:
        """Create business logic embeddings with smart deduplication"""
        stats = {"inserted": 0, "updated": 0, "skipped": 0, "reused": 0, "relinked": 0}
        
        for rule_idx, rule_text in enumerate(rules):
            if not rule_text.strip():
                continue
                
            try:
                content_hash = cls.generate_content_hash(rule_text)
                # Check if this exact content already exists
                existing_by_hash = await cls.check_embedding_by_content_hash(
                    connection, user_id, project_id, content_hash, "business_logic_embeddings"
                )
                
                if existing_by_hash:
                     # Case A: content identical, but maybe belongs to another document
                    if existing_by_hash.get("document_id") != document_id:
                        await cls._update_business_logic_document_id(
                            connection,
                            existing_by_hash["id"],
                            document_id
                        )
                        stats["relinked"] += 1
                        logger.info(f"Relinked existing embedding (rule {rule_idx + 1}) to new document ID {document_id}")
                    else:
                        stats["reused"] += 1
                        logger.info(f"Reused existing business logic embedding for rule {rule_idx + 1}")
                        continue
                
                # Check if embedding exists for this specific document+rule
                existing_for_document = await cls._check_existing_business_logic_embedding(
                    connection, document_id, user_id, project_id, rule_idx + 1
                )
                
                # Generate embedding only when needed
                new_embedding = await get_embedding_func(rule_text)
                
                metadata = {
                    "rule_number": rule_idx + 1,
                    "rule_type": "business_rule",
                    "word_count": len(rule_text.split()),
                    "content_hash": content_hash
                }
                
                if existing_for_document:
                    # Update existing embedding
                    await cls._update_business_logic_embedding(
                        connection, document_id, user_id, project_id,
                        rule_idx + 1, rule_text, new_embedding, metadata
                    )
                    stats["updated"] += 1
                    logger.info(f"Updated business logic embedding for rule {rule_idx + 1}")
                else:
                    # Insert new embedding
                    await cls._insert_business_logic_embedding(
                        connection, document_id, user_id, project_id,
                        rule_idx + 1, rule_text, new_embedding, metadata
                    )
                    stats["inserted"] += 1
                    logger.info(f"Created new business logic embedding for rule {rule_idx + 1}")
                    
            except Exception as e:
                logger.error(f"Error processing business rule {rule_idx}: {e}")
                continue
                
                # new_embedding = await get_embedding_func(rule_text)
                # should_update, action = cls.should_update_embedding(
                #     existing, rule_text, new_embedding
                # )
                
            #     if should_update:
            #         metadata = {
            #             "rule_number": rule_idx + 1,
            #             "rule_type": "business_rule",
            #             "word_count": len(rule_text.split()),
            #             "content_hash": cls.generate_content_hash(rule_text)
            #         }
                    
            #         if action == "INSERT":
            #             await cls._insert_business_logic_embedding(
            #                 connection, document_id, user_id, project_id,
            #                 rule_idx + 1, rule_text, new_embedding, metadata
            #             )
            #             stats["inserted"] += 1
            #         else:
            #             await cls._update_business_logic_embedding(
            #                 connection, document_id, user_id, project_id,
            #                 rule_idx + 1, rule_text, new_embedding, metadata
            #             )
            #             stats["updated"] += 1
            #     else:
            #         stats["skipped"] += 1
                    
            # except Exception as e:
            #     logger.error(f"Error processing business rule {rule_idx}: {e}")
            #     continue
        
        return stats

    @staticmethod
    async def _check_existing_business_logic_embedding(
        connection, document_id: str, user_id: int, project_id: str, rule_number: int
    ) -> Optional[Dict[str, Any]]:
        """Check if business logic embedding already exists for this rule"""
        query = """
        SELECT content, embedding, metadata 
        FROM business_logic_embeddings 
        WHERE user_id = $1 AND project_id = $2 AND rule_number = $3
        """
        
        try:
            row = await connection.fetchrow(query, user_id, project_id, rule_number)
            return dict(row) if row else None
        except Exception as e:
            logger.error(f"Error checking existing business logic embedding: {e}")
            return None
    
    @staticmethod
    async def _insert_business_logic_embedding(
        connection, document_id, user_id, project_id,
        rule_number, content, embedding, metadata
    ):
        """Insert new business logic embedding"""
        insert_query = """
        INSERT INTO business_logic_embeddings
        (document_id, project_id, user_id, rule_number, content, embedding, metadata)
        VALUES ($1, $2, $3, $4, $5, $6::vector, $7)
        """
        # Debug logging
        logger.debug(f"Inserting business logic embedding: rule={rule_number}")
        logger.debug(f"Embedding type: {type(embedding)}, length: {len(embedding) if isinstance(embedding, list) else 'N/A'}")
        await connection.execute(
            insert_query, document_id, project_id, user_id,
            rule_number, content, embedding, json.dumps(metadata)
        )
    
    @staticmethod
    async def _update_business_logic_embedding(
        connection, document_id, user_id, project_id,
        rule_number, content, embedding, metadata
    ):
        """Update existing business logic embedding"""
        update_query = """
        UPDATE business_logic_embeddings 
        SET content = $1, embedding = $2::vector, metadata = $3, document_id = $4, created_at = CURRENT_TIMESTAMP
        WHERE user_id = $5 AND project_id = $6 AND rule_number = $7
        """
        # Debug logging
        logger.debug(f"Updating business logic embedding: rule={rule_number}")
        logger.debug(f"Embedding type: {type(embedding)}, length: {len(embedding) if isinstance(embedding, list) else 'N/A'}")
        await connection.execute(
            update_query, content, embedding, json.dumps(metadata),
            document_id, user_id, project_id, rule_number
        )
    @staticmethod
    async def _update_business_logic_document_id(connection, row_id: str, new_document_id: str):
        """Update the document_id for an existing embedding record"""
        query = """
        UPDATE business_logic_embeddings
        SET document_id = $1, created_at = CURRENT_TIMESTAMP
        WHERE id = $2
        """
        await connection.execute(query, new_document_id, row_id)


class ReferenceParser(DocumentParser):
    """Parser for reference documents with hash-based duplicate checking"""
    
    @staticmethod
    def split_text_into_chunks(text: str, max_chunk_size: int = 1000) -> List[str]:
        """Split text into manageable chunks"""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk + paragraph) < max_chunk_size:
                current_chunk += paragraph + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

    @classmethod
    async def create_embeddings_with_dedup(
        cls,
        connection,
        chunks: List[str],
        document_id: str,
        user_id: int,
        project_id: str,
        get_embedding_func
    ) -> Dict[str, int]:
        """Create reference embeddings with smart deduplication"""
        stats = {"inserted": 0, "updated": 0, "skipped": 0, "reused": 0, "relinked": 0}
        
        for chunk_idx, chunk_text in enumerate(chunks):
            if not chunk_text.strip():
                continue
                    
            try:
                content_hash = cls.generate_content_hash(chunk_text)
                
                # Check if this exact content already exists
                existing_by_hash = await cls.check_embedding_by_content_hash(
                    connection, user_id, project_id, content_hash, "reference_embeddings"
                )
                
                if existing_by_hash:
                    # Case A: Content exists but belongs to a different document - RELINK
                    if existing_by_hash.get("document_id") != document_id:
                        await cls._update_reference_document_id(
                            connection,
                            existing_by_hash["id"],
                            document_id
                        )
                        stats["relinked"] += 1
                        logger.info(f"Relinked existing embedding (chunk {chunk_idx}) to new document ID {document_id}")
                    else:
                        # Case B: Content exists and already belongs to this document - REUSE
                        stats["reused"] += 1
                        logger.info(f"Reused existing reference embedding for chunk {chunk_idx}")
                    continue
                
                # Check if embedding exists for this specific document+chunk position
                existing_for_document = await cls._check_existing_reference_embedding(
                    connection, user_id, project_id, chunk_idx
                )
                
                # Generate embedding only when needed (not found by hash)
                new_embedding = await get_embedding_func(chunk_text)
                
                metadata = {
                    "chunk_index": chunk_idx,
                    "chunk_type": "reference_content",
                    "word_count": len(chunk_text.split()),
                    "content_hash": content_hash
                }
                
                if existing_for_document:
                    # Case C: Different content at the same chunk position - UPDATE
                    await cls._update_reference_embedding(
                        connection, document_id, user_id, project_id,
                        chunk_idx, chunk_text, new_embedding, metadata
                    )
                    stats["updated"] += 1
                    logger.info(f"Updated reference embedding for chunk {chunk_idx}")
                else:
                    # Case D: Completely new content at new position - INSERT
                    await cls._insert_reference_embedding(
                        connection, document_id, user_id, project_id,
                        chunk_idx, chunk_text, new_embedding, metadata
                    )
                    stats["inserted"] += 1
                    logger.info(f"Created new reference embedding for chunk {chunk_idx}")
                    
            except Exception as e:
                logger.error(f"Error processing reference chunk {chunk_idx}: {e}")
                continue
        
        return stats

    @staticmethod
    async def _check_existing_reference_embedding(
        connection, user_id: int, project_id: str, chunk_index: int
    ) -> Optional[Dict[str, Any]]:
        """Check if reference embedding already exists for this chunk"""
        query = """
        SELECT content, embedding, metadata 
        FROM reference_embeddings 
        WHERE user_id = $1 AND project_id = $2 AND chunk_index = $3
        """
        
        try:
            row = await connection.fetchrow(query, user_id, project_id, chunk_index)
            return dict(row) if row else None
        except Exception as e:
            logger.error(f"Error checking existing reference embedding: {e}")
            return None

    @staticmethod
    async def _insert_reference_embedding(
        connection, document_id: str, user_id: int, project_id: str,
        chunk_index: int, content: str, embedding: List[float], metadata: Dict
    ):
        """Insert new reference embedding"""
        insert_query = """
        INSERT INTO reference_embeddings
        (document_id, project_id, user_id, chunk_index, content, embedding, metadata)
        VALUES ($1, $2, $3, $4, $5, $6::vector, $7)
        """
        # Debug logging
        logger.debug(f"Inserting reference embedding: chunk={chunk_index}")
        logger.debug(f"Embedding type: {type(embedding)}, length: {len(embedding) if isinstance(embedding, list) else 'N/A'}")
        await connection.execute(
            insert_query, document_id, project_id, user_id,
            chunk_index, content, embedding, json.dumps(metadata)
        )

    @staticmethod
    async def _update_reference_embedding(
        connection, document_id: str, user_id: int, project_id: str,
        chunk_index: int, content: str, embedding: List[float], metadata: Dict
    ):
        """Update existing reference embedding"""
        update_query = """
        UPDATE reference_embeddings 
        SET content = $1, embedding = $2::vector, metadata = $3, document_id = $4, created_at = CURRENT_TIMESTAMP
        WHERE user_id = $5 AND project_id = $6 AND chunk_index = $7
        """
        # Debug logging
        logger.debug(f"Updating reference embedding: chunk={chunk_index}")
        logger.debug(f"Embedding type: {type(embedding)}, length: {len(embedding) if isinstance(embedding, list) else 'N/A'}")
        await connection.execute(
            update_query, content, embedding, json.dumps(metadata),
            document_id, user_id, project_id, chunk_index
        )
    @staticmethod
    async def _update_reference_document_id(connection, row_id: str, new_document_id: str):
        """Update the document_id for an existing embedding record - RELINK operation"""
        query = """
        UPDATE reference_embeddings
        SET document_id = $1, created_at = CURRENT_TIMESTAMP
        WHERE id = $2
        """
        await connection.execute(query, new_document_id, row_id)

    @classmethod
    async def batch_insert_reference_embeddings(
        cls,
        connection,
        embeddings_batch: List[Dict[str, Any]]
    ) -> int:
        """Batch insert multiple reference embeddings for efficiency"""
        if not embeddings_batch:
            return 0
            
        insert_query = """
        INSERT INTO reference_embeddings 
        (document_id, project_id, user_id, chunk_index, content, embedding, metadata)
        VALUES ($1, $2, $3, $4, $5, $6::vector, $7)
        ON CONFLICT (document_id, user_id, project_id, chunk_index) 
        DO UPDATE SET 
            content = EXCLUDED.content,
            embedding = EXCLUDED.embedding,
            metadata = EXCLUDED.metadata,
            created_at = CURRENT_TIMESTAMP
        """
        
        batch_values = []
        for emb in embeddings_batch:
            batch_values.append((
                emb['document_id'],
                emb['project_id'],
                emb['user_id'],
                emb['chunk_index'],
                emb['content'],
                emb['embedding'],
                json.dumps(emb['metadata'])
            ))
        
        await connection.executemany(insert_query, batch_values)
        return len(batch_values)

class FileExtractor:
    """Extract text from different file formats"""
    
    @staticmethod
    async def extract_text(file_content: bytes, file_path: str) -> str:
        """Extract text from various file formats"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return FileExtractor._extract_pdf_text(file_content)
            elif file_extension in ['.docx', '.doc']:
                return FileExtractor._extract_docx_text(file_content)
            elif file_extension in ['.txt', '.json']:
                return file_content.decode('utf-8')
            else:
                return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return ""
    
    @staticmethod
    def _extract_pdf_text(file_content: bytes) -> str:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    @staticmethod
    def _extract_docx_text(file_content: bytes) -> str:
        # """Extract text from DOCX"""
        # doc = Document(BytesIO(file_content))
        # text = ""
        # for paragraph in doc.paragraphs:
        #     text += paragraph.text + "\n"
        # return text
        """Extract text from DOCX — includes numbered list values"""
        doc = Document(BytesIO(file_content))
        text_lines = []
        list_counter = 1

        for para in doc.paragraphs:
            p_text = para.text.strip()
            if not p_text:
                continue

            # Detect if this paragraph is part of a numbered/bulleted list
            if para._p.pPr is not None and para._p.pPr.numPr is not None:
                # It’s part of a numbered list → prepend visible number
                text_lines.append(f"{list_counter}. {p_text}")
                list_counter += 1
            else:
                text_lines.append(p_text)

        return "\n".join(text_lines)
